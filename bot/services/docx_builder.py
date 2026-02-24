import os
from datetime import datetime

from docx import Document
from docx.shared import Pt

from bot.config import TEMP_DIR


def _setup_style(doc: Document) -> None:
    """Настроить стиль документа."""
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = "Arial"


def _add_text_page(doc: Document, text: str) -> None:
    """Добавить текст как обычный абзац."""
    doc.add_paragraph(text)


def _add_table_page(doc: Document, table_data: list[list[str]]) -> None:
    """Добавить таблицу в документ."""
    if not table_data or not table_data[0]:
        return
    rows = len(table_data)
    cols = len(table_data[0])

    table = doc.add_table(rows=rows, cols=cols, style="Table Grid")

    for ri, row in enumerate(table_data):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = cell_text
            # Размер шрифта в ячейках чуть меньше
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def _create_document(pages: list[dict]) -> Document:
    """Создать документ Word из списка страниц.

    Каждая страница — dict:
        {"type": "text", "data": "текст"}
        или
        {"type": "table", "data": [[cell, ...], ...]}
    """
    doc = Document()
    _setup_style(doc)

    for i, page in enumerate(pages):
        if page["type"] == "table":
            _add_table_page(doc, page["data"])
        else:
            _add_text_page(doc, page["data"])
        # Разрыв страницы после каждой, кроме последней
        if i < len(pages) - 1:
            doc.add_page_break()

    return doc


def build_single_docx(pages: list[dict]) -> str:
    """Все страницы в один .docx файл. Возвращает путь к файлу."""
    os.makedirs(TEMP_DIR, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filepath = os.path.join(TEMP_DIR, f"ocr_{timestamp}.docx")

    doc = _create_document(pages)
    doc.save(filepath)
    return filepath


def build_separate_docx(pages: list[dict]) -> list[str]:
    """Каждая страница в отдельный .docx файл. Возвращает список путей."""
    os.makedirs(TEMP_DIR, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    paths = []

    for i, page in enumerate(pages, start=1):
        filepath = os.path.join(TEMP_DIR, f"ocr_{timestamp}_{i}.docx")
        doc = _create_document([page])
        doc.save(filepath)
        paths.append(filepath)

    return paths
